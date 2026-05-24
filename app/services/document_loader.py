"""Policy document loading and chunking utilities."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile
import re
import xml.etree.ElementTree as ET


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass(frozen=True)
class Document:
    source: str
    text: str


@dataclass(frozen=True)
class DocumentChunk:
    chunk_id: str
    source: str
    text: str


def load_documents(source_dir: Path) -> list[Document]:
    """Load supported policy documents from a directory tree."""
    documents: list[Document] = []
    if not source_dir.exists():
        return documents

    for path in sorted(source_dir.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        text = extract_text(path)
        if text:
            documents.append(Document(source=str(path.relative_to(source_dir)), text=text))
    return documents


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".txt", ".md"}:
        return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    return ""


def chunk_documents(documents: list[Document], chunk_size: int = 900, overlap: int = 160) -> list[DocumentChunk]:
    """Split documents into overlapping text chunks."""
    chunks: list[DocumentChunk] = []
    for document in documents:
        text = _normalize_text(document.text)
        if not text:
            continue
        start = 0
        index = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"{document.source}#{index}",
                        source=document.source,
                        text=chunk_text,
                    )
                )
            if end >= len(text):
                break
            start = max(0, end - overlap)
            index += 1
    return chunks


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF extraction requires the 'pypdf' package.") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _normalize_text("\n".join(pages))


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return _extract_docx_with_zip(path)

    document = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return _normalize_text("\n".join(paragraphs))


def _extract_docx_with_zip(path: Path) -> str:
    """Fallback DOCX text extraction using only the standard library."""
    with ZipFile(path) as archive:
        xml_bytes = archive.read("word/document.xml")

    root = ET.fromstring(xml_bytes)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    texts = [node.text or "" for node in root.findall(".//w:t", namespace)]
    return _normalize_text("\n".join(texts))


def _normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
